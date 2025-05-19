from docx import Document

def read_doc():
    doc = Document('app/data/Dense Medium Separation of Coal.docx')
    res = [
    ]
    for para in doc.paragraphs:
        if para.text == '':
            continue
        res.append({
            'style': para.style.name,
            'text': para.text
        })
    return res

def read_report(report):
    if report == 'Progress Report':
        doc = Document('app/data/Dense Medium Separation of Coal.docx')
    elif report == 'Thesis Report':
        doc = Document('app/data/THESIS REPORT MAIN.docx')

    res = []
    for para in doc.paragraphs:
        if para.text == '':
            continue
        res.append({
            'style': para.style.name,
            'text': para.text
        })
    return res



if __name__ == '__main__':
    read_doc()